"""
Text and DOCX parsers for documentation ingestion.

Handles parsing of plain text and DOCX files with structure detection.
"""

import asyncio
import logging
import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument

from navigator.knowledge.ingest.documentation.metadata import caption_image_with_vision

logger = logging.getLogger(__name__)


def parse_plain_text(file_path: Path) -> tuple[str, dict[str, Any]]:
	"""
	Parse plain text file with structure detection.
	
	Detects sections, lists, and other structure patterns.
	
	Args:
		file_path: Path to text file
	
	Returns:
		Tuple of (file content, metadata dict)
	"""
	with open(file_path, 'r', encoding='utf-8') as f:
		content = f.read()

	metadata = {}

	# Detect structure patterns
	lines = content.split('\n')
	sections = sum(1 for line in lines if line.strip().startswith('#') or (line.strip() and line.strip().isupper() and len(line.strip()) < 100))
	lists = sum(1 for line in lines if line.strip().startswith('-') or line.strip().startswith('*') or line.strip().startswith('•'))
	numbered_lists = sum(1 for line in lines if re.match(r'^\d+[\.\)]\s+', line.strip()))

	metadata['sections'] = sections
	metadata['lists'] = lists
	metadata['numbered_lists'] = numbered_lists

	return content, metadata


def parse_docx(file_path: Path, image_caption_cache: dict[str, str] | None = None) -> tuple[str, dict[str, Any]]:
	"""
	Parse DOCX file with comprehensive feature extraction.
	
	Extracts text, tables, images, metadata, and structure.
	
	Args:
		file_path: Path to DOCX file
		image_caption_cache: Optional cache for image captions
	
	Returns:
		Tuple of (extracted content, metadata dict)
	"""
	if image_caption_cache is None:
		image_caption_cache = {}

	content_parts = []
	metadata = {}

	try:
		doc = DocxDocument(file_path)

		# Extract document metadata
		core_props = doc.core_properties
		metadata['author'] = core_props.author or ''
		metadata['title'] = core_props.title or ''
		metadata['subject'] = core_props.subject or ''
		metadata['keywords'] = core_props.keywords or ''
		metadata['category'] = core_props.category or ''
		metadata['comments'] = core_props.comments or ''
		metadata['creation_date'] = str(core_props.created) if core_props.created else ''
		metadata['modification_date'] = str(core_props.modified) if core_props.modified else ''

		# Extract content
		tables_found = []
		images_found = []
		hyperlinks_found = []

		for para_idx, paragraph in enumerate(doc.paragraphs):
			text = paragraph.text.strip()

			if not text:
				continue

			# Check paragraph style (heading, list, etc.)
			style_name = paragraph.style.name if paragraph.style else ''

			# Headings
			if 'Heading' in style_name:
				level = 1
				if 'Heading 1' in style_name:
					level = 1
				elif 'Heading 2' in style_name:
					level = 2
				elif 'Heading 3' in style_name:
					level = 3
				elif 'Heading 4' in style_name:
					level = 4
				elif 'Heading 5' in style_name:
					level = 5
				elif 'Heading 6' in style_name:
					level = 6
				content_parts.append(f"\n{'#' * level} {text}\n")

			# Lists
			elif 'List' in style_name or paragraph.style.name.startswith('List'):
				content_parts.append(f"- {text}\n")

			# Regular paragraphs
			else:
				# Check for hyperlinks in paragraph
				for run in paragraph.runs:
					if run.hyperlink:
						href = run.hyperlink.target if run.hyperlink.target else ''
						link_text = run.text
						hyperlinks_found.append({'text': link_text, 'href': href})
						content_parts.append(f"[{link_text}]({href})")
					else:
						content_parts.append(run.text)
				content_parts.append("\n")

		# Extract tables
		for table_idx, table in enumerate(doc.tables):
			table_data = []
			content_parts.append(f"\n[Table {table_idx + 1} Start]\n")

			headers = []
			for row_idx, row in enumerate(table.rows):
				cells = [cell.text.strip() for cell in row.cells]

				if row_idx == 0:
					headers = cells
					content_parts.append(" | ".join(cells) + "\n")
					content_parts.append(" | ".join(['---'] * len(cells)) + "\n")
				else:
					content_parts.append(" | ".join(cells) + "\n")
					if headers:
						table_data.append(dict(zip(headers, cells)))

			content_parts.append(f"[Table {table_idx + 1} End]\n")
			if table_data:
				tables_found.append(table_data)

		# Extract images and generate vision captions
		image_count = 0
		image_refs = []

		# Extract images from document relationships
		for rel in doc.part.rels.values():
			if "image" in rel.target_ref:
				image_count += 1

				try:
					# Get image data from relationship
					image_part = rel.target_part
					image_bytes = image_part.blob

					# Determine image format from relationship
					image_ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else "png"
					image_format = image_ext.lower() if image_ext.lower() in ['png', 'jpeg', 'jpg', 'gif'] else "png"

					# Try to generate vision caption (if vision APIs available)
					# Note: This is async, so we'll do it synchronously using asyncio.run if no event loop
					try:
						loop = asyncio.get_event_loop()
						if loop.is_running():
							# Event loop running, can't use asyncio.run - skip vision for now
							caption = None
						else:
							caption = asyncio.run(caption_image_with_vision(image_bytes, image_format, image_caption_cache=image_caption_cache))
					except RuntimeError:
						# No event loop, create new one
						caption = asyncio.run(caption_image_with_vision(image_bytes, image_format, image_caption_cache=image_caption_cache))

					# Add image reference with caption if available
					image_info = {
						'format': image_format,
						'size': len(image_bytes)
					}
					if caption:
						image_info['caption'] = caption
						# Embed caption in content stream
						content_parts.append(f"\n[Image: {caption}]\n")
					else:
						content_parts.append("\n[Image: (no alt text or caption available)]\n")

					image_refs.append(image_info)
				except Exception as e:
					logger.debug(f"Failed to extract/caption image from DOCX: {e}")
					# Still count the image even if extraction fails
					image_refs.append({'format': 'unknown', 'error': str(e)})

		metadata['tables'] = len(tables_found)
		metadata['images'] = image_count
		metadata['hyperlinks'] = len(hyperlinks_found)
		metadata['table_data'] = tables_found
		metadata['hyperlink_refs'] = hyperlinks_found[:50]  # First 50
		metadata['image_refs'] = image_refs  # Include image references with captions

	except Exception as e:
		logger.error(f"Error parsing DOCX {file_path}: {e}", exc_info=True)
		raise

	return "".join(content_parts), metadata
